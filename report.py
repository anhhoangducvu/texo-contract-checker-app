# -*- coding: utf-8 -*-
"""
Sinh báo cáo rà soát hợp đồng ra file Word (.docx) — khổ A4, Times New Roman,
header/footer chuyên nghiệp, bảng màu theo mức rủi ro, khối chữ ký Phòng Kỹ thuật.
Báo cáo ĐỘC LẬP (không tham chiếu hợp đồng/dự án khác).

Đơn vị tư vấn luôn là TEXO; đơn vị rà soát là Phòng Kỹ thuật; Trưởng phòng: Hoàng Đức Vũ.
Các thông tin này cố định (không đổi qua UI) theo yêu cầu.
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from knowledge import DO, CAM, XANH, LEVEL_LABEL

# ----- Nhận diện đơn vị (cố định) -----
COMPANY_NAME = "CÔNG TY CỔ PHẦN TEXO TƯ VẤN VÀ ĐẦU TƯ"
DEPARTMENT = "PHÒNG KỸ THUẬT"
HEAD_OF_DEPT = "Hoàng Đức Vũ"          # Trưởng phòng Kỹ thuật (ký báo cáo)

# ----- Màu -----
C_DO = RGBColor(0xC0, 0x00, 0x00)
C_CAM = RGBColor(0xC4, 0x59, 0x11)
C_XANH = RGBColor(0x2E, 0x7D, 0x32)
C_HEAD = RGBColor(0x1F, 0x37, 0x64)
C_GREY = RGBColor(0x70, 0x70, 0x70)
LEVEL_RGB = {DO: C_DO, CAM: C_CAM, XANH: C_XANH}
LEVEL_FILL = {DO: "F8CBAD", CAM: "FCE4D6", XANH: "E2EFDA"}
LEVEL_ICON = {DO: "🔴", CAM: "🟠", XANH: "🟢"}
FLAG_ICON = "🚩"

FONT = "Times New Roman"
CONTENT_WIDTH = 9026  # DXA, A4 trừ lề 1 inch mỗi bên


# --------------------------------------------------------------------------- #
# Tiện ích định dạng
# --------------------------------------------------------------------------- #
def _set_font(run, size=13, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:cs"), FONT)


def _shade(cell, fill_hex):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcpr.append(shd)


def _para_border(p, edge="bottom", color="2E5496", sz=6, space=4):
    ppr = p._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), str(space))
    e.set(qn("w:color"), color)
    pbdr.append(e)


def _para(doc, text="", size=13, bold=False, color=None, italic=False,
          align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        r = p.add_run(text)
        _set_font(r, size, bold, color, italic)
    return p


def _heading(doc, text, size=14):
    return _para(doc, text, size=size, bold=True, color=C_HEAD, space_after=6)


def _cell_text(cell, text, size=11, bold=False, color=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    _set_font(r, size, bold, color, italic)
    return p


def _add_cell_line(cell, label, text, size=11):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(label)
    _set_font(r1, size, bold=True)
    r2 = p.add_run(text)
    _set_font(r2, size)


def _page_field(paragraph):
    run = paragraph.add_run()
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
    run._r.append(fc1); run._r.append(it); run._r.append(fc2)
    _set_font(run, 9, color=C_GREY)


def _cell_border(cell, edge, color, sz):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), color)
    borders.append(e)


def _add_tab_run(paragraph, color=None, size=9):
    r = paragraph.add_run()
    r._r.append(OxmlElement("w:tab"))
    _set_font(r, size, color=color)


def _setup_page(doc):
    sec = doc.sections[0]
    sec.page_width = Twips(11906)
    sec.page_height = Twips(16838)
    sec.left_margin = Twips(1440)
    sec.right_margin = Twips(1440)
    sec.top_margin = Twips(1260)
    sec.bottom_margin = Twips(1080)
    sec.header_distance = Twips(560)
    sec.footer_distance = Twips(560)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(13)
    _decorate_section(sec)
    return sec


def _decorate_section(section):
    # ----- Header (letterhead) chạy mọi trang: 2 cột, công ty trái – phòng phải, gạch dưới
    hdr = section.header
    for para in list(hdr.paragraphs):
        para._p.getparent().remove(para._p)
    tbl = hdr.add_table(rows=1, cols=2, width=Twips(CONTENT_WIDTH))
    tbl.autofit = False
    left, right = tbl.rows[0].cells
    left.width = Twips(6200)
    right.width = Twips(CONTENT_WIDTH - 6200)
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(3)
    _set_font(lp.add_run(COMPANY_NAME), 9, bold=True, color=C_HEAD)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(3)
    _set_font(rp.add_run(DEPARTMENT.title()), 9, italic=True, color=C_GREY)
    for cell in (left, right):
        _cell_border(cell, "bottom", "2E5496", 6)
    endp = hdr.add_paragraph()
    endp.paragraph_format.space_after = Pt(0)
    endp.paragraph_format.space_before = Pt(0)
    _set_font(endp.add_run(""), 2)

    # ----- Footer chạy mọi trang -----
    fp = section.footer.paragraphs[0]
    fp.text = ""
    fp.paragraph_format.tab_stops.add_tab_stop(Twips(CONTENT_WIDTH), WD_TAB_ALIGNMENT.RIGHT)
    rf = fp.add_run(f"Phòng Kỹ thuật — Trưởng phòng: {HEAD_OF_DEPT}")
    _set_font(rf, 9, color=C_GREY)
    _add_tab_run(fp, color=C_GREY)
    rt = fp.add_run("Trang ")
    _set_font(rt, 9, color=C_GREY)
    _page_field(fp)
    _para_border(fp, "top", "BFBFBF", 4, 4)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _signature_block(doc):
    now = datetime.now()
    _para(doc, "", space_after=6)
    _para(doc, f"Hà Nội, ngày {now:%d} tháng {now:%m} năm {now:%Y}",
          size=12, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = t.rows[0].cells

    def col(cell, title, name):
        cell.text = ""
        p0 = cell.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(0)
        _set_font(p0.add_run(title), 12, bold=True)
        p1 = cell.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p1.add_run("(Ký, ghi rõ họ tên)"), 11, italic=True)
        for _ in range(3):
            cell.add_paragraph()
        p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p2.add_run(name), 13, bold=True)

    col(cells[0], "NGƯỜI RÀ SOÁT", "")
    col(cells[1], "TRƯỞNG PHÒNG KỸ THUẬT", HEAD_OF_DEPT)
    for c in cells:
        c.width = Twips(CONTENT_WIDTH // 2)
    _remove_table_borders(t)


def _title_block(doc, subtitle, extra_line=""):
    # Tên công ty/phòng đã nằm ở header letterhead -> tiêu đề KHÔNG lặp lại, vào thẳng tên báo cáo.
    _para(doc, "", space_after=6)
    _para(doc, "BÁO CÁO RÀ SOÁT PHÁP LÝ HỢP ĐỒNG", size=18, bold=True, color=C_HEAD,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    _para(doc, subtitle, size=12, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    if extra_line:
        _para(doc, extra_line, size=10, italic=True, color=C_GREY,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    rule = _para(doc, "", space_after=10)
    _para_border(rule, "bottom", "2E5496", 8, 1)


# =========================================================================== #
# BÁO CÁO RÚT GỌN (rule-based)
# =========================================================================== #
def build_report(result: dict) -> bytes:
    doc = Document()
    _setup_page(doc)

    ctx = result["context"]
    meta = result["meta"]
    summary = result["summary"]
    roles = result.get("roles") or {}
    party = roles.get("party_label") or "Đơn vị tư vấn (Bên B)"

    _title_block(
        doc,
        f"(Góc nhìn {party[0].upper() + party[1:]} — Bên B)",
        f"Ngày lập: {datetime.now().strftime('%d/%m/%Y')}  •  "
        f"Đơn vị rà soát: {DEPARTMENT.title()} — Công cụ rà soát tự động (không dùng AI)")

    # 1. Thông tin chung
    _heading(doc, "1. THÔNG TIN CHUNG & BỐI CẢNH")
    _para(doc, f"• Tệp hợp đồng: {meta['filename']}", space_after=2)
    _para(doc, f"• Quy mô: {meta['n_paragraphs']} đoạn, {meta['n_headings']} đề mục, "
               f"{meta['n_comments']} ghi chú (comment).", space_after=2)
    _para(doc, f"• Kiểu kết cấu: {ctx['kieu_ket_cau']}.", space_after=2)
    langs = ", ".join(ctx["content_languages"]) or "vi"
    _para(doc, f"• Ngôn ngữ: {langs}"
               f"{' (song ngữ)' if ctx['song_ngu'] else ''}.", space_after=2)
    _para(doc, f"• Nguồn vốn: {ctx['nguon_von']}.", space_after=2)
    _para(doc, f"• Trần phạt áp dụng: {ctx['tran_phat']}.", space_after=2)
    _para(doc, f"• Đối chiếu mẫu TT 02/2023/TT-BXD: {ctx['ghi_chu_tt02']}", space_after=8)

    # 1b. Vai trò
    pm = roles.get("primary_meta")
    if pm:
        _heading(doc, "2. VAI TRÒ TƯ VẤN & PHẠM VI CHUẨN")
        _para(doc, f"• Vai trò chính (đoán): {pm['name']}.", space_after=2)
        if roles.get("multi_role"):
            others = ", ".join(m["name"] for m in roles["present_meta"] if m["id"] != pm["id"])
            _para(doc, f"• Có thể GỘP NHIỀU vai trò (cũng thấy: {others}) → tách phạm vi & căn cứ.",
                  color=C_CAM, space_after=2)
        _para(doc, f"• Phạm vi chuẩn: {pm['scope']}", size=11, space_after=2)
        _para(doc, f"• Căn cứ pháp lý: {pm['basis']}", size=11, space_after=2)
        if pm.get("independence"):
            _para(doc, f"• Yêu cầu độc lập: {pm['independence']}", size=11, space_after=2)
        for note in roles.get("cross_role_flags", []):
            _para(doc, f"{FLAG_ICON} Vượt vai trò: {note}", size=11, bold=True, color=C_DO, space_after=2)
        _para(doc, "", space_after=6)
        n3, n4, n5 = "3", "4", "5"
    else:
        n3, n4, n5 = "2", "3", "4"

    # Tóm tắt
    _heading(doc, f"{n3}. TỔNG QUAN CHUNG")
    _para(doc,
          f"Phát hiện {summary[DO]} rủi ro mức CAO (đỏ), {summary[CAM]} rủi ro mức "
          f"TRUNG BÌNH (cam), và {summary['missing']} chủ đề chuẩn không thấy xuất hiện "
          f"trong hợp đồng (trong đó {summary['missing_do']} thuộc nhóm trọng yếu).",
          space_after=4)
    _para(doc,
          "Lưu ý: kết quả rà soát SƠ BỘ bằng công cụ tự động (đối chiếu mẫu câu và độ phủ "
          "chủ đề), KHÔNG thay thế ý kiến luật sư. Mọi phát hiện cần người có chuyên môn "
          "kiểm chứng lại trên bản hợp đồng gốc trước khi đàm phán/ký.",
          size=11, italic=True, space_after=8)

    # Bảng rủi ro
    _heading(doc, f"{n4}. BẢNG PHÂN TÍCH RỦI RO PHÁT HIỆN")
    findings = result["findings"]
    if not findings:
        _para(doc, "Không phát hiện mẫu câu rủi ro điển hình nào. Vẫn nên rà soát thủ công "
                   "và xem mục độ phủ chủ đề bên dưới.", italic=True, space_after=8)
    else:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        widths = [1500, 4400, 3126]
        hdr = table.rows[0].cells
        for c, label in enumerate(["Mức / Chủ đề", "Vấn đề & căn cứ", "Đề xuất đàm phán"]):
            _cell_text(hdr[c], label, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            _shade(hdr[c], "1F3764")
        for f in findings:
            row = table.add_row().cells
            _cell_text(row[0], f"{LEVEL_ICON[f['level']]} {LEVEL_LABEL[f['level']]}", size=11, bold=True,
                       color=LEVEL_RGB[f["level"]])
            _add_cell_line(row[0], "Chủ đề ", str(f["topic"]))
            p = row[0].add_paragraph()
            r = p.add_run(f["label"]); _set_font(r, 10, italic=True)
            _shade(row[0], LEVEL_FILL[f["level"]])
            _cell_text(row[1], f["problem"], size=11)
            _add_cell_line(row[1], "Trích: ", f"“{f['quote']}”", size=10)
            _add_cell_line(row[1], "Căn cứ: ", f["basis"], size=10)
            _cell_text(row[2], f["suggest"], size=11)
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Twips(w)
        _para(doc, "", space_after=6)

    # Độ phủ
    _heading(doc, f"{n5}. ĐỘ PHỦ 21 CHỦ ĐỀ CHUẨN (chủ đề THIẾU cũng là rủi ro)")
    missing = [c for c in result["coverage"] if not c["present"]]
    if not missing:
        _para(doc, "Hợp đồng có nhắc tới toàn bộ 21 chủ đề chuẩn.", italic=True, space_after=8)
    else:
        t2 = doc.add_table(rows=1, cols=3)
        t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        t2.style = "Table Grid"
        w2 = [800, 3200, 5026]
        hdr = t2.rows[0].cells
        for c, label in enumerate(["Mức", "Chủ đề thiếu", "Khuyến nghị"]):
            _cell_text(hdr[c], label, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            _shade(hdr[c], "1F3764")
        for m in missing:
            row = t2.add_row().cells
            _cell_text(row[0], f"{LEVEL_ICON[m['missing_risk']]} {LEVEL_LABEL[m['missing_risk']]}",
                       size=10, bold=True, color=LEVEL_RGB[m["missing_risk"]])
            _cell_text(row[1], f"{m['id']}. {m['name']}", size=11)
            _cell_text(row[2], m["missing_note"], size=11)
            _shade(row[0], LEVEL_FILL[m["missing_risk"]])
        for row in t2.rows:
            for i, w in enumerate(w2):
                row.cells[i].width = Twips(w)
        _para(doc, "", space_after=6)

    if result["comments"]:
        _heading(doc, "GHI CHÚ (COMMENT) CÓ SẴN TRONG FILE")
        for c in result["comments"]:
            _para(doc, f"• [{c['author']} – {c['date']}] {c['text']}", size=11, space_after=2)
        _para(doc, "", space_after=6)

    _heading(doc, "LƯU Ý SỬ DỤNG")
    _para(doc,
          "Báo cáo do công cụ tự động lập theo bộ tiêu chí rà soát hợp đồng tư vấn xây dựng "
          "của TEXO, không sử dụng AI. Công cụ phát hiện theo MẪU CÂU và TỪ KHÓA nên có thể "
          "bỏ sót cách diễn đạt lạ hoặc báo nhầm khi câu chữ trùng từ khóa. Hãy coi đây là "
          "bước sàng lọc đầu tiên; quyết định đàm phán/ký kết phải dựa trên rà soát của "
          "người có chuyên môn trên bản gốc.", size=11, space_after=4)

    _signature_block(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =========================================================================== #
# BÁO CÁO ĐẦY ĐỦ DO AI LẬP
# =========================================================================== #
_AI_LEVEL = {"ĐỎ": DO, "CAM": CAM, "XANH": XANH, "DO": DO, "RED": DO,
             "CAO": DO, "TRUNG BÌNH": CAM, "TRUNG BINH": CAM}


def _norm_level(v):
    if not v:
        return CAM
    return _AI_LEVEL.get(str(v).strip().upper(), CAM)


def build_ai_report(result: dict, ai: dict) -> bytes:
    data = ai.get("data", {}) if ai else {}
    doc = Document()
    _setup_page(doc)

    meta = result.get("meta", {})
    roles = result.get("roles") or {}
    party = roles.get("party_label") or "Đơn vị tư vấn (Bên B)"

    _title_block(
        doc,
        f"(Góc nhìn {party[0].upper() + party[1:]} — Bên B) — BẢN ĐẦY ĐỦ",
        f"Ngày lập: {datetime.now().strftime('%d/%m/%Y')}  •  "
        f"Đơn vị rà soát: {DEPARTMENT.title()}  •  "
        f"Hỗ trợ phân tích bởi AI: {ai.get('provider','')} / {ai.get('model','')}")

    # ---- BẢNG THÔNG TIN HỢP ĐỒNG ----
    hd = data.get("thong_tin_hop_dong") or {}
    _heading(doc, "1. THÔNG TIN HỢP ĐỒNG")
    meta_rows = [
        ("Tên dự án / Công trình", hd.get("ten_du_an") or meta.get("filename", "")),
        ("Chủ đầu tư (Bên A)",     hd.get("ten_cdt", "—")),
        ("Đơn vị tư vấn (Bên B)",  hd.get("ten_tu_van", "—")),
        ("Số hiệu hợp đồng",        hd.get("so_hop_dong", "—")),
        ("Gói thầu",                hd.get("goi_thau", "—")),
        ("Giá trị hợp đồng",        hd.get("gia_hop_dong", "—")),
        ("Thời hạn thực hiện",      hd.get("thoi_han", "—")),
        ("Loại hợp đồng",           hd.get("loai_hop_dong", "—")),
    ]
    tbl_hd = doc.add_table(rows=len(meta_rows), cols=2)
    tbl_hd.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_hd.style = "Table Grid"
    for i, (label, value) in enumerate(meta_rows):
        cells = tbl_hd.rows[i].cells
        _cell_text(cells[0], label, size=11, bold=True)
        _shade(cells[0], "D9E1F2")
        _cell_text(cells[1], value, size=11)
        cells[0].width = Twips(2800)
        cells[1].width = Twips(6226)
    _para(doc, "", space_after=6)

    _heading(doc, "2. BỐI CẢNH & VAI TRÒ")
    if data.get("vai_tro"):
        _para(doc, f"• Vai trò tư vấn: {data['vai_tro']}", space_after=2)
    if data.get("nguon_von"):
        _para(doc, f"• Nguồn vốn & trần phạt: {data['nguon_von']}", space_after=2)
    for n in roles.get("cross_role_flags", []):
        _para(doc, f"{FLAG_ICON} Vượt vai trò: {n}", bold=True, color=C_DO, space_after=2)
    _para(doc, "", space_after=4)

    _heading(doc, "3. TỔNG QUAN CHUNG")
    tqc = (data.get("tong_quan_chung") or data.get("tom_tat_dieu_hanh")
           or "(AI không trả về tổng quan.)")
    _para(doc, tqc, space_after=8)

    _heading(doc, "4. PHÂN TÍCH CHI TIẾT — 21 CHỦ ĐỀ RỦI RO CHUẨN")
    rows = data.get("dieu_khoan") or []
    if not rows:
        _para(doc, "(AI không trả về bảng điều khoản.)", italic=True, space_after=8)
    else:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        widths = [1600, 4300, 3126]
        hdr = table.rows[0].cells
        for c, label in enumerate(["Chủ đề / Điều / Mức", "Vấn đề & căn cứ pháp lý", "Đề xuất đàm phán"]):
            _cell_text(hdr[c], label, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            _shade(hdr[c], "1F3764")
        for it in rows:
            lvl = _norm_level(it.get("muc_do"))
            row = table.add_row().cells
            # Cột 1: Chủ đề + tên điều
            chu_de = str(it.get("chu_de", ""))
            ref_txt = str(it.get("ref", ""))
            _cell_text(row[0], f"{LEVEL_ICON[lvl]} {LEVEL_LABEL[lvl]}", size=11, bold=True, color=LEVEL_RGB[lvl])
            if chu_de:
                p_cd = row[0].add_paragraph()
                r_cd = p_cd.add_run(f"Chủ đề {chu_de}"); _set_font(r_cd, 10, bold=True)
            if ref_txt and ref_txt != "Thiếu":
                p_ref = row[0].add_paragraph()
                r_ref = p_ref.add_run(ref_txt); _set_font(r_ref, 10, italic=True)
            elif ref_txt == "Thiếu":
                p_ref = row[0].add_paragraph()
                r_ref = p_ref.add_run("⚠ Thiếu điều khoản"); _set_font(r_ref, 10, bold=True)
                r_ref.font.color.rgb = C_CAM
            _shade(row[0], LEVEL_FILL[lvl])
            _cell_text(row[1], str(it.get("van_de", "")), size=11)
            _cell_text(row[2], str(it.get("de_xuat", "")), size=11)
        for row in table.rows:
            for i, w in enumerate(widths):
                            row.cells[i].width = Twips(w)
        _para(doc, "", space_after=6)

    favs = data.get("dieu_khoan_co_loi") or []
    if favs:
        _heading(doc, "5. ĐIỀU KHOẢN CÓ LỢI CẦN BẢO VỆ")
        for x in favs:
            _para(doc, f"• {x}", size=12, space_after=2)
        _para(doc, "", space_after=4)

    miss = data.get("thieu_sot") or []
    if miss:
        _heading(doc, "6. NỘI DUNG CÒN THIẾU CẦN BỔ SUNG")
        for x in miss:
            _para(doc, f"• {x}", size=12, space_after=2)
        _para(doc, "", space_after=4)

    pri = data.get("uu_tien_dam_phan") or {}
    if pri:
        _heading(doc, "7. THỨ TỰ ƯU TIÊN ĐÀM PHÁN")
        for k, title, col in [("phai_dat", "Phải đạt", C_DO),
                              ("nen_dat", "Nên đạt", C_CAM),
                              ("don_dep", "Dọn dẹp câu chữ", C_XANH)]:
            items = pri.get(k) or []
            if items:
                _para(doc, title + ":", bold=True, color=col, space_after=2)
                for x in items:
                    _para(doc, f"• {x}", size=12, space_after=1)
        _para(doc, "", space_after=4)

    _heading(doc, "8. LƯU Ý SỬ DỤNG")
    _para(doc,
          "Báo cáo bản đầy đủ do mô hình AI hỗ trợ phân tích trên nền bộ tiêu chí rà soát "
          "hợp đồng tư vấn của TEXO, kết hợp lớp quét tự động làm điểm tựa. AI có thể sai "
          "sót hoặc diễn giải chưa chuẩn; mọi nhận định cần người có chuyên môn kiểm chứng "
          "trên bản hợp đồng gốc trước khi đàm phán/ký. Đây không phải ý kiến luật sư chính thức.",
          size=11, space_after=4)

    _signature_block(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
